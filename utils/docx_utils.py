from docx import Document
from docx.shared import Pt
from datetime import datetime
import re
import pythoncom
import comtypes.client
import os

def create_letter(template_path, student_name, parent_name, address, branch, attendance, roll_no):
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    doc = Document(template_path)
    current_date = datetime.now().strftime('%Y-%m-%d')

    replacements = {
        '{CURRENT_DATE}': current_date,
        '{STUDENT_NAME}': student_name,
        '{PARENT_NAME}': parent_name,
        '{BRANCH}': branch,
        '{ADDRESS}': address,
        '{ATTENDANCE}': str(attendance),
        '{ROLL_NUMBER}': str(roll_no)
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)

    letter_filename = f"generated_letters/{roll_no}_letter_{timestamp}.docx"
    doc.save(letter_filename)

    pdf_filename = f"generated_letters/{roll_no}_letter_{timestamp}.pdf"
    convert_docx_to_pdf(letter_filename, pdf_filename)

    return pdf_filename

def convert_docx_to_pdf(docx_filename, pdf_filename):
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(docx_filename))
    doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)
    doc.Close()
    word.Quit()