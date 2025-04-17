# routes/health.py

from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import re
import uuid

bp = Blueprint('health', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

# Predefined templates
predefined_templates = {
    "Medical Check-up Camps": """
Dear Students,

We are pleased to inform you that a {{TYPE}} Health Check-up Camp has been scheduled on {{DATE}} at the Health Center.

This initiative is aimed at promoting student wellness and early detection of health issues. All students are encouraged to take advantage of this opportunity and get themselves examined by certified medical professionals.

Please make sure to carry your ID card and report at the venue on time.

Stay healthy and take care.

""",

    "Mental Health & Counseling Sessions": """
Dear Students,

We understand the importance of mental well-being and are happy to announce that one-on-one counseling sessions are now available through the Student Mentorship Program.

If you feel the need to talk or seek professional support, we encourage you to reach out to your assigned mentor or contact the Counseling Cell for assistance.

Your privacy will be respected, and support will be offered in a safe and confidential environment.

Take care of yourself. You are not alone.

""",

    "Blood Donation Drives": """
Dear Students,

A Blood Donation Camp has been organized on {{DATE}} in collaboration with the NSS Unit.

All willing donors are requested to participate in this noble cause. Your contribution can help save lives and support emergency medical needs.

Certificates will be provided to all donors. Light refreshments will also be arranged post-donation.

Let’s come together for a greater good. Donate blood, save lives.

""",

    "Vaccination Drives": """
Dear Students,

We are organizing a {{TYPE}} Vaccination Drive scheduled on {{DATE}} at the Medical Room.

All students who have not yet received this vaccination are advised to participate. This is a preventive health measure to ensure the well-being of the campus community.

Please carry your health records (if any) and student ID card. The process will be supervised by trained medical personnel.

Your health is our priority. Stay protected.

""",

    "Health Guidelines & Quarantine Policies": """
Dear Students,

This is to inform you that the Health Office has issued revised health and safety guidelines effective immediately.

As per the new protocol, Please continue to maintain basic hygiene practices and monitor your health regularly. The full guideline document is available on the College Website.

Let us work together to maintain a safe and healthy campus environment.

"""
}


def format_date(date_str):
    dt = datetime.strptime(date_str, "%Y-%m-%dT%H:%M")
    day = dt.day
    suffix = 'th' if 4 <= day <= 20 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return dt.strftime(f"%d{suffix} %B %Y %I:%M %p")

@bp.route('/generate-health-notice', methods=['POST'])
def generate_notice():
    data = request.get_json()
    notif_type = data.get('type')
    date = data.get('date', '')
    kind = data.get('kind', '')  # for check-up/vaccine
    subject = data.get('subject') or notif_type
    content = data.get('content', '')

    current_date = datetime.now().strftime('%B %d, %Y')
    formatted_date = format_date(date) if date else ''
    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').replace('/', '_').replace('&', '_').lower()}_{unique_id}"
    word_path = f"{GENERATED_LETTERS_PATH}/{filename_base}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    doc = Document("templates/Custom.docx")

    if notif_type in predefined_templates:
        temp = predefined_templates[notif_type]
        content = temp.replace("{{TYPE}}", kind).replace("{{DATE}}", formatted_date)

    replacements = {
        "{CURRENT_DATE}": current_date,
        "{SOURCE}": "PRINCIPAL",
        "{SUBJECT}": subject,
        "{CONTENT}": content,
        "{DATE}": formatted_date,
        "{COPY_TO}": "All HODs\nHealth Center\nStudents"
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(word_path)

    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Health notice generated successfully!',
        'download_link': f'http://localhost:5001/download/{filename_base}.pdf'
    })

@bp.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(GENERATED_LETTERS_PATH, filename, as_attachment=True)