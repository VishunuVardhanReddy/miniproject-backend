# routes/fee_letter.py
from flask import Blueprint, request, jsonify
from database import get_db_connection
import os, re, base64, sqlite3
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import pythoncom
import comtypes.client
from PyPDF2 import PdfMerger

bp = Blueprint('fee_letter', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

@bp.route('/fee-letter', methods=['POST'])
def fee_letter():
    file = request.files['file']
    due_date_str = request.form['dueDate']
    today = datetime.now().date()
    selected_due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()

    if selected_due_date <= today:
        return jsonify({"error": "Due date must be later than today's date."}), 400

    file_path = os.path.join("uploads", file.filename)
    file.save(file_path)

    df = pd.read_excel(file_path)
    numeric_columns = [
        '1st year payable', '1ST YEAR paid', 
        '2nd year payable', '2ND YEAR paid', 
        '3rd year payable', '3RD YEAR paid', 
        '4th year payable', '4TH YEAR paid'
    ]
    df[numeric_columns] = df[numeric_columns].apply(pd.to_numeric, errors='coerce')

    conn = get_db_connection()
    cursor = conn.cursor()
    fee_letters = []

    for _, row in df.iterrows():
        roll_no = row['Roll_no']
        year = row['Year']
        amount_due = 0

        if year == 1:
            amount_due = row['1st year payable'] - row['1ST YEAR paid']
        elif year == 2:
            amount_due = sum([
                row['1st year payable'] - row['1ST YEAR paid'],
                row['2nd year payable'] - row['2ND YEAR paid']
            ])
        elif year == 3:
            amount_due = sum([
                row['1st year payable'] - row['1ST YEAR paid'],
                row['2nd year payable'] - row['2ND YEAR paid'],
                row['3rd year payable'] - row['3RD YEAR paid']
            ])
        elif year == 4:
            amount_due = sum([
                row['1st year payable'] - row['1ST YEAR paid'],
                row['2nd year payable'] - row['2ND YEAR paid'],
                row['3rd year payable'] - row['3RD YEAR paid'],
                row['4th year payable'] - row['4TH YEAR paid']
            ])

        if amount_due > 0:
            cursor.execute("SELECT student_name, parent_name, address, branch FROM students WHERE Roll_no=?", (roll_no,))
            student = cursor.fetchone()
            if student:
                student_name, parent_name, address, branch = student
                letter_filename = create_fee_letter(
                    "templates/Fees.docx", student_name, parent_name, address, branch,
                    roll_no, amount_due, selected_due_date, year, row
                )
                fee_letters.append(letter_filename)

    conn.close()

    pdf_files = [convert_docx_to_pdf(file) for file in fee_letters]
    combined_pdf_filename = merge_pdfs(pdf_files)

    with open(combined_pdf_filename, 'rb') as f:
        pdf_binary_content = f.read()

    pdf_base64 = base64.b64encode(pdf_binary_content).decode('utf-8')

    return jsonify({
        "message": "Fee letter generation successful!",
        "generated_letters": {
            "count": len(fee_letters),
            "pdf_preview": pdf_base64
        }
    })

def create_fee_letter(template_path, student_name, parent_name, address, branch, roll_no, payable_amount, due_date, year, row):
    current_date = datetime.now().strftime('%Y-%m-%d')
    due_date_str = due_date.strftime('%Y-%m-%d')
    doc = Document(template_path)

    replacements = {
        '{CURRENT_DATE}': current_date,
        '{STUDENT_NAME}': student_name,
        '{PARENT_NAME}': parent_name,
        '{BRANCH}': branch,
        '{ADDRESS}': address,
        '{DUE_DATE}': due_date_str,
        '{PAYABLE_AMOUNT}': str(payable_amount),
        '{ROLL_NUMBER}': str(roll_no)
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        if table.rows[0].cells[0].text == "YEAR":
            if year >= 1:
                add_fee_row(table, "1st Year", row['1st year payable'], row['1ST YEAR paid'])
            if year >= 2:
                add_fee_row(table, "2nd Year", row['2nd year payable'], row['2ND YEAR paid'])
            if year >= 3:
                add_fee_row(table, "3rd Year", row['3rd year payable'], row['3RD YEAR paid'])
            if year >= 4:
                add_fee_row(table, "4th Year", row['4th year payable'], row['4TH YEAR paid'])

    letter_filename = f"generated_letters/{roll_no}_fee_letter_{current_date}.docx"
    doc.save(letter_filename)
    return letter_filename

def add_fee_row(table, year_text, payable, paid):
    row = table.add_row().cells
    row[0].text = year_text
    row[1].text = str(payable)
    row[2].text = str(paid)
    row[3].text = str(payable - paid)

def convert_docx_to_pdf(docx_filename):
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    pdf_filename = docx_filename.replace('.docx', '.pdf')
    doc = word.Documents.Open(os.path.abspath(docx_filename))
    doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)
    doc.Close()
    word.Quit()
    return pdf_filename

def merge_pdfs(pdf_files):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    output_path = f"generated_letters/combined_fee_letters_{timestamp}.pdf"
    merger.write(output_path)
    merger.close()
    return output_path