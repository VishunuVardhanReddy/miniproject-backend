from flask import Blueprint, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import sqlite3
import os
import re
import uuid

bp = Blueprint('academic_misconduct', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

def fetch_student_details(roll_no):
    conn = sqlite3.connect('students_database_unique_names.db')
    cursor = conn.cursor()
    cursor.execute("SELECT parent_name, address, student_name FROM students WHERE Roll_no = ?", (roll_no,))
    result = cursor.fetchone()
    conn.close()
    return result

@bp.route('/generate-academic-misconduct', methods=['POST'])
def generate_academic_misconduct():
    data = request.get_json()
    roll_no = data.get('rollNo')
    meeting_date = data.get('meetingDate')
    reason = data.get('reason')
    description = data.get('description') or ''

    if description:
        description = f" - {description}"

    student_data = fetch_student_details(roll_no)
    if not student_data:
        return jsonify({'error': 'Student not found'}), 404

    parent_name, address, student_name = student_data
    current_date = datetime.now().strftime('%B %d, %Y')

    unique_id = uuid.uuid4().hex
    word_path = f"{GENERATED_LETTERS_PATH}/academic_{roll_no}_{unique_id}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    doc = Document('templates/AcademicMisconduct.docx')

    replacements = {
        '{PARENT_NAME}': parent_name,
        '{ADDRESS}': address,
        '{STUDENT_NAME}': student_name,
        '{CURRENT_DATE}': current_date,
        '{REASON}': reason,
        '{DESCRIPTION}': description,
        '{ROLL_NO}': roll_no,
        '{MEETING_DATE}': meeting_date
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(word_path)

    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Academic Misconduct letter generated!',
        'download_link': f'http://localhost:5001/download/academic_{roll_no}_{unique_id}.pdf'
    })