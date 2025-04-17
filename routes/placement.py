from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import re
import uuid

bp = Blueprint('placement', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

# Predefined templates
predefined_templates = {
    "Company Shortlists": """
Dear Students,

We are pleased to inform you that the shortlist for the upcoming interview round conducted by {{COMPANY}} has been released.

Students are advised to check the Training & Placement portal for their names or reach out to Placement Officer for further instructions.

Please prepare accordingly and stay updated for additional communication.

Best of luck!
""",
    "Pre-Placement Talks (PPTs)": """
Dear Students,

You are invited to attend a Pre-Placement Talk by {{COMPANY}}, scheduled on {{DATE}}. 

The session will be conducted online and provide an overview of the company, roles, and career growth opportunities.

Attendance is highly recommended for all eligible students.
""",
    "Internship Opportunities": """
Dear Students,

We are excited to announce a summer internship opportunity with {{COMPANY}}.

Interested students are encouraged to apply before {{DATE}}.

This internship is a stepping stone to industry exposure and future opportunities.
""",
    "Offer Letter / Joining Updates": """
Dear Students,

We are happy to announce that offer letters from {{COMPANY}} have been issued.

Students reach out to Placement Officer for further instructions.

Congratulations and best wishes!
"""
}

# Helper function to format the date in the desired format
def format_drive_date(drive_date):
    # Parse the ISO format string (e.g., "2025-04-24T10:30")
    dt = datetime.strptime(drive_date, "%Y-%m-%dT%H:%M")

    # Get the day with the correct suffix
    day = dt.day
    if 4 <= day <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    # Format the date as "24th April 2025 10:30 AM"
    formatted_date = dt.strftime(f"%d{suffix} %B %Y %I:%M %p")
    return formatted_date

@bp.route('/generate-placement-notice', methods=['POST'])
def generate_notice():
    data = request.get_json()
    notif_type = data.get('type')
    company = data.get('company', '')
    date = data.get('date', '')
    subject = data.get('subject', notif_type)  # Default subject is the notification type
    content = data.get('content', '')
    venue = data.get('venue', 'Lab 2')  # Default venue for workshops

    # If no subject is provided, use the notification type as subject
    if not subject:
        subject = notif_type

    current_date = datetime.now().strftime('%B %d, %Y')
    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').replace('/', '_').replace('&', '_').lower()}_{unique_id}"
    word_path = f"{GENERATED_LETTERS_PATH}/{filename_base}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    # Format the DRIVE_DATE field
    formatted_drive_date = format_drive_date(date)

    # Template selection
    if notif_type == "Upcoming Placement Drives":
        doc = Document("templates/Placement.docx")
        replacements = {
            "{CURRENT_DATE}": current_date,
            "{COMPANY}": company,
            "{DRIVE_DATE}": formatted_drive_date  # Use the formatted drive date here
        }
    else:
        doc = Document("templates/Custom.docx")
        if notif_type in predefined_templates:
            temp_content = predefined_templates[notif_type]
            content = temp_content.replace("{{COMPANY}}", company).replace("{{DATE}}", formatted_drive_date).replace("{{VENUE}}", venue)
        
        replacements = {
            "{CURRENT_DATE}": current_date,
            "{SOURCE}": "PLACEMENT CELL",
            "{SUBJECT}": subject,  # Use the populated subject here
            "{CONTENT}": content,
            "{DATE}": formatted_drive_date,
            "{COPY_TO}": "HOD\nIndividual\nR&D"  # Replacing {COPY_TO} with new lines
        }

    # Replace in document paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    # Save DOCX
    doc.save(word_path)

    # Convert DOCX to PDF using Word COM
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Notice generated successfully!',
        'download_link': f'http://localhost:5001/download/{filename_base}.pdf'
    })