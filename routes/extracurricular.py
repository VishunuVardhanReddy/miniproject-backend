from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import uuid

bp = Blueprint('extracurricular', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

predefined_templates = {
    "Inter-College Fest Invites": """
Dear Students,

You are invited to participate in the upcoming Inter-College Fest scheduled on {{DATE}} at {{PLACE}}.

It is a celebration of talent, creativity, and collaboration among institutions. Interested students are requested to register before {{DEADLINE}} to ensure their participation.

Let’s represent our college with pride!
""",
    "Competition Alerts": """
Dear Students,

We are excited to announce an upcoming competition: {{COMPETITION_NAME}}, which will be held on {{DATE}} at {{PLACE}}.

All students interested in participating are required to register before {{DEADLINE}}. This is your chance to showcase your skills and bring glory to our institution!

Don't miss out—get involved and shine!
""",
    "Achievements / Highlights": """
Dear Students,

We are proud to share the recent achievements of our students in various events and competitions:

{{ACHIEVEMENTS}}

Let’s congratulate them for their efforts and excellence!

Keep up the great work, everyone!
"""
}

def format_date(date_str, with_time=False):
    if not date_str:
        return ""
    fmt = "%Y-%m-%dT%H:%M" if "T" in date_str else "%Y-%m-%d"
    dt = datetime.strptime(date_str, fmt)
    return dt.strftime("%d %B %Y %I:%M %p") if with_time else dt.strftime("%d %B %Y")

@bp.route('/generate-extracurricular-notice', methods=['POST'])
def generate_notice():
    data = request.get_json()
    notif_type = data.get('type')

    if notif_type in predefined_templates:
        subject = notif_type
    else:
        subject = data.get('subject', notif_type)

    content = ""
    template_file = "Custom.docx"
    current_date = datetime.now().strftime('%d %B %Y')

    replacements = {
        "{SUBJECT}": subject
    }
    doc = None

    if notif_type == "Inter-College Fest Invites":
        content = predefined_templates[notif_type] \
            .replace("{{DATE}}", format_date(data.get("date"))) \
            .replace("{{PLACE}}", data.get("place", "")) \
            .replace("{{DEADLINE}}", format_date(data.get("deadline")))
        replacements.update({
            "{CONTENT}": content,
            "{SUBJECT}": subject
        })
        doc = Document(f"templates/Custom.docx")

    elif notif_type == "Competition Alerts":
        content = predefined_templates[notif_type] \
            .replace("{{DATE}}", format_date(data.get("date"))) \
            .replace("{{PLACE}}", data.get("place", "")) \
            .replace("{{DEADLINE}}", format_date(data.get("deadline"))) \
            .replace("{{COMPETITION_NAME}}", data.get("competition_name", ""))
        replacements.update({
            "{CONTENT}": content,
            "{SUBJECT}": subject
        })
        doc = Document(f"templates/Custom.docx")

    elif notif_type == "Achievements / Highlights":
        achievements = data.get("achievements", [])
        formatted = "\n".join([f"• {a}" for a in achievements])
        content = predefined_templates[notif_type].replace("{{ACHIEVEMENTS}}", formatted)
        replacements.update({
            "{CONTENT}": content,
            "{SUBJECT}": subject
        })
        doc = Document(f"templates/Custom.docx")

    elif notif_type == "Event Announcements":
        template_file = "Festival.docx"
        replacements.update({
            "{FESTIVAL}": data.get("festival", ""),
            "{DATE}": format_date(data.get("date")),
            "{FACULTY}": data.get("faculty", ""),
            "{POST}": data.get("post", ""),
            "{BRANCH}": data.get("branch", ""),
            "{EVENT}": data.get("event_name", ""),
            "{TIME}": data.get("time", "")
        })
        doc = Document(f"templates/{template_file}")

    elif notif_type == "NSS Event":
        template_file = "NSS.docx"
        replacements.update({
            "{EVENT_DATE}": format_date(data.get("event_date")),
            "{PLACE}": data.get("place", ""),
            "{DEADLINE}": format_date(data.get("deadline"))
        })
        doc = Document(f"templates/{template_file}")

    elif notif_type == "Workshop":
        template_file = "Workshop.docx"
        replacements.update({
            "{TOPIC}": data.get("topic", ""),
            "{DATE}": format_date(data.get("date")),
            "{TIME}": format_date(data.get("time")),
            "{PLACE}": data.get("place", "")
        })
        doc = Document(f"templates/{template_file}")

    elif notif_type == "Girls Sports":
        template_file = "girls_sport.docx"
        doc = Document(f"templates/{template_file}")
        replacements.update({
            "{SPORT}": data.get("sport", ""),
            "{STARTDATE}": format_date(data.get("start_date")),
            "{ENDDATE}": format_date(data.get("end_date")),
            "{PLACE}": data.get("place", "")
        })

        # Replace placeholders in paragraphs
        for p in doc.paragraphs:
            for run in p.runs:
                for key, val in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Fill the table dynamically
        schedule = data.get("schedule", [])
        for table in doc.tables:
            if schedule and len(table.rows) > 1:
                placeholder_row = table.rows[1]
                for idx, row_data in enumerate(schedule, start=1):
                    new_row = table.add_row()
                    sr = str(idx)
                    date = format_date(row_data.get("date", ""))
                    branch = row_data.get("branch", "")

                    new_row.cells[0].text = sr
                    new_row.cells[1].text = date
                    new_row.cells[2].text = branch

                    for i in range(3):
                        for p in new_row.cells[i].paragraphs:
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                table._tbl.remove(placeholder_row._tr)

    else:
        content = data.get("content", "")
        subject = data.get("subject", "Custom Notice")
        replacements = {
            "{SUBJECT}": subject,
            "{CONTENT}": content
        }
        doc = Document(f"templates/Custom.docx")
        for p in doc.paragraphs:
            for run in p.runs:
                for key, val in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for key, val in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

    # Common replacements
    replacements["{CURRENT_DATE}"] = current_date
    replacements["{SOURCE}"] = "PRINCIPAL"
    replacements["{COPY_TO}"] = "All Departments HODs\nDean Office\nR&D\nStudents"

    # Generate unique file paths
    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').replace('/', '_').replace('&', '_').lower()}_{unique_id}"
    word_path = f"{GENERATED_LETTERS_PATH}/{filename_base}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    if doc:
        for p in doc.paragraphs:
            for run in p.runs:
                for key, val in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for key, val in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

        doc.save(word_path)

        pythoncom.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        docx_file = word.Documents.Open(os.path.abspath(word_path))
        docx_file.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        docx_file.Close()
        word.Quit()

    return jsonify({
        'message': 'Notice generated successfully!',
        'download_link': f'http://localhost:5001/download/{filename_base}.pdf'
    })