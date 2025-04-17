from flask import Blueprint, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import sqlite3
import os
import re
import uuid

bp = Blueprint('warning', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

# 🔹 Helper: Fetch student info from DB
def fetch_student_details(roll_no):
    conn = sqlite3.connect('students_database_unique_names.db')
    cursor = conn.cursor()
    cursor.execute("SELECT parent_name, address, student_name FROM students WHERE Roll_no = ?", (roll_no,))
    result = cursor.fetchone()
    conn.close()
    return result

@bp.route('/generate-warning', methods=['POST'])
def generate_warning():
    data = request.get_json()
    roll_no = data.get('rollNo')
    mistake_date = data.get('mistakeDate')
    reason = data.get('reason')
    description = data.get('description') or ''

    # Add hyphen around description if it is provided
    if description:
        description = f" - {description}"

    student_data = fetch_student_details(roll_no)
    if not student_data:
        return jsonify({'error': 'Student not found'}), 404

    parent_name, address, student_name = student_data
    current_date = datetime.now().strftime('%B %d, %Y')  # e.g. April 15, 2025

    # Unique file naming
    unique_id = uuid.uuid4().hex
    word_path = f"{GENERATED_LETTERS_PATH}/warning_{roll_no}_{unique_id}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    # Load template
    doc = Document('templates/warningnotice.docx')

    # Placeholder replacements
    replacements = {
        '{PARENT_NAME}': parent_name,
        '{ADDRESS}': address,
        '{STUDENT_NAME}': student_name,
        '{CURRENT_DATE}': current_date,
        '{REASON}': reason,
        '{DESCRIPTION}': description,
        '{DATE}': mistake_date,
        '{ROLL_NUMBER}': roll_no
    }

    # 🔸 Replace in paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # 🔸 Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    # Save DOCX
    doc.save(word_path)

    # Convert DOCX ➝ PDF using Word (COM automation)
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Letter generated successfully!',
        'download_link': f'http://localhost:5001/download/warning_{roll_no}_{unique_id}.pdf'
    })

@bp.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(GENERATED_LETTERS_PATH, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True, mimetype='application/pdf')
    return jsonify({'error': 'File not found'}), 404