from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import re
import uuid

bp = Blueprint('administrative', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

def format_date(date_str):
    try:
        return datetime.strptime(date_str, "%Y-%m-%dT%H:%M").strftime("%d %B %Y")
    except ValueError:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d %B %Y")

@bp.route('/generate-admin-notice', methods=['POST'])
def generate_notice():
    data = request.get_json()
    notif_type = data.get('type')
    subject = data.get('subject', notif_type)
    content = data.get('content', '')
    template_file = ""
    placeholders = {}

    current_date = datetime.now().strftime('%B %d, %Y')
    formatted_date = lambda d: format_date(d) if d else ''

    if notif_type == "Re-Admission":
        template_file = "templates/ReAdmission.docx"
        placeholders = {
            "{NAME}": data.get('student_name', ''),
            "{ROLL_NO}": data.get('roll_no', ''),
            "{BRANCH}": data.get('branch', ''),
            "{YEAR_SEMESTER}": data.get('year_semester', ''),
            "{RE_DATE}": formatted_date(data.get('date', ''))
        }

    elif notif_type == "Academic Regulation Change":
        template_file = "templates/Regulation.docx"
        placeholders = {
            "{OLD_REGULATION}": data.get('old_reg', ''),
            "{NEW_REGULATION}": data.get('new_reg', ''),
            "{EFFECTIVE_DATE}": formatted_date(data.get('date', ''))
        }

    elif notif_type == "Rules & Regulation Updates":
        template_file = "templates/Rules.docx"
        placeholders = {
            "{EFFECTIVE_DATE}": formatted_date(data.get('date', ''))
        }

    elif notif_type == "Staff Updates":
        template_file = "templates/Staff_updates.docx"
        placeholders = {}

    else:  # Others
        template_file = "templates/Custom.docx"
        placeholders = {
            "{SUBJECT}": subject,
            "{CONTENT}": content
        }

    # Common replacements
    placeholders["{CURRENT_DATE}"] = current_date
    placeholders["{SOURCE}"] = "ADMINISTRATION"
    placeholders["{COPY_TO}"] = "All Departments\nDean Office\nStudents"

    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').lower()}_{unique_id}"
    word_path = f"{GENERATED_LETTERS_PATH}/{filename_base}.docx"
    pdf_path = word_path.replace(".docx", ".pdf")

    doc = Document(template_file)

    for p in doc.paragraphs:
        for run in p.runs:
            for key, val in placeholders.items():
                if key in run.text:
                    run.text = re.sub(re.escape(key), val, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in placeholders.items():
                            if key in run.text:
                                run.text = re.sub(re.escape(key), val, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(word_path)

    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Notice generated successfully!',
        'download_link': f'http://localhost:5001/download/{filename_base}.pdf'
    })

@bp.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(GENERATED_LETTERS_PATH, filename, as_attachment=True)