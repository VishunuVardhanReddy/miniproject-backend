# routes/hostel_fee_letter.py

from flask import Blueprint, request, jsonify
from database import get_db_connection
import os, re, base64
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt
import pythoncom
import comtypes.client
from PyPDF2 import PdfMerger

bp = Blueprint('hostel_fee_letter', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

@bp.route('/hostel-fee-letter', methods=['POST'])
def hostel_fee_letter():
    file = request.files['file']
    due_date_str = request.form['dueDate']
    today = datetime.now().date()
    selected_due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()

    if selected_due_date <= today:
        return jsonify({"error": "Due date must be later than today's date."}), 400

    file_path = os.path.join("uploads", file.filename)
    file.save(file_path)

    df = pd.read_excel(file_path)
    numeric_columns = ['1ST YEAR paid', '2ND YEAR paid', '3RD YEAR paid', '4TH YEAR paid']
    df[numeric_columns] = df[numeric_columns].apply(pd.to_numeric, errors='coerce')

    conn = get_db_connection()
    cursor = conn.cursor()
    fee_letters = []

    for _, row in df.iterrows():
        roll_no = row['Roll_no']
        year = row['Year']
        total_fees = row['Fees']

        due_1st_year = total_fees - row['1ST YEAR paid'] if year >= 1 else 0
        due_2nd_year = total_fees - row['2ND YEAR paid'] if year >= 2 else 0
        due_3rd_year = total_fees - row['3RD YEAR paid'] if year >= 3 else 0
        due_4th_year = total_fees - row['4TH YEAR paid'] if year >= 4 else 0
        total_due = due_1st_year + due_2nd_year + due_3rd_year + due_4th_year

        if total_due > 0:
            cursor.execute("SELECT student_name, parent_name, address, branch FROM students WHERE Roll_no=?", (roll_no,))
            student = cursor.fetchone()

            if student:
                student_name, parent_name, address, branch = student
                letter_filename = create_fee_letter(
                    'templates/Hostel-fee.docx', student_name, parent_name, address, branch, roll_no,
                    due_1st_year, due_2nd_year, due_3rd_year, due_4th_year, selected_due_date, year, row
                )
                fee_letters.append(letter_filename)

    conn.close()

    if not fee_letters:
        return jsonify({"message": "No letters were generated."}), 200

    pdf_files = [convert_docx_to_pdf(f) for f in fee_letters]
    combined_pdf_filename = merge_pdfs(pdf_files)

    with open(combined_pdf_filename, 'rb') as f:
        pdf_base64 = base64.b64encode(f.read()).decode('utf-8')

    return jsonify({
        "message": "Fee letter generation successful!",
        "generated_letters": {
            "count": len(fee_letters),
            "pdf_preview": pdf_base64
        }
    })

def create_fee_letter(template_path, student_name, parent_name, address, branch, roll_no,
                      due_1st_year, due_2nd_year, due_3rd_year, due_4th_year, due_date, year, row):
    current_date = datetime.now().strftime('%Y-%m-%d')
    due_date_str = due_date.strftime('%Y-%m-%d')
    doc = Document(template_path)

    payable_amount = due_1st_year + due_2nd_year + due_3rd_year + due_4th_year
    replacements = {
        '{CURRENT_DATE}': current_date,
        '{STUDENT_NAME}': student_name,
        '{PARENT_NAME}': parent_name,
        '{BRANCH}': branch,
        '{ADDRESS}': address,
        '{DUE_DATE}': due_date_str,
        '{ROLL_NUMBER}': str(roll_no),
        '{PAYABLE_AMOUNT}': str(payable_amount)
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        if table.rows[0].cells[0].text == "YEAR":
            if year >= 1:
                add_row(table, "1st Year", row['Fees'], row['1ST YEAR paid'], due_1st_year)
            if year >= 2:
                add_row(table, "2nd Year", row['Fees'], row['2ND YEAR paid'], due_2nd_year)
            if year >= 3:
                add_row(table, "3rd Year", row['Fees'], row['3RD YEAR paid'], due_3rd_year)
            if year >= 4:
                add_row(table, "4th Year", row['Fees'], row['4TH YEAR paid'], due_4th_year)

    filename = f"{GENERATED_LETTERS_PATH}/{roll_no}_hostel_fee_letter_{current_date}.docx"
    doc.save(filename)
    return filename

def add_row(table, year_text, total_fee, paid, due):
    row = table.add_row().cells
    row[0].text = year_text
    row[1].text = str(total_fee)
    row[2].text = str(paid)
    row[3].text = str(due)

def convert_docx_to_pdf(docx_filename):
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    pdf_filename = docx_filename.replace('.docx', '.pdf')
    doc = word.Documents.Open(os.path.abspath(docx_filename))
    doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)
    doc.Close()
    word.Quit()
    return pdf_filename

def merge_pdfs(pdf_files):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    output_path = f"{GENERATED_LETTERS_PATH}/combined_hostel_fee_letters_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.pdf"
    merger.write(output_path)
    merger.close()
    return output_path