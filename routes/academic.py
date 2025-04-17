from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import uuid

bp = Blueprint('academic', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

# Templates for content-based notices
predefined_templates = {
    "Class Schedule Changes": """
Dear Students,

Please be informed that there are changes to the class schedule. The new schedule will begin from {{START_DATE}}. Ensure that you check the updated timetable for your respective classes.

Thank you for your attention!
""",
    "Exam Notifications": """
Dear Students,

We would like to inform you about the upcoming exams. The exams are scheduled as follows:

Type: {{EXAM_TYPE}}  
Date: {{EXAM_DATE}}

Please prepare accordingly.

Best of luck!
""",
    "Result Announcements": """
Dear Students,

We are pleased to announce that the results will be declared on {{RESULT_DATE}}. Please check the College Official Portal for your results.

Best regards,  
The Examination Committee
"""
}

def format_date(date_str, with_time=False):
    if not date_str:
        return ""
    fmt = "%Y-%m-%dT%H:%M" if "T" in date_str else "%Y-%m-%d"
    dt = datetime.strptime(date_str, fmt)
    return dt.strftime("%d %B %Y %I:%M %p") if with_time else dt.strftime("%d %B %Y")

@bp.route('/generate-academic-notice', methods=['POST'])
def generate_academic_notice():
    data = request.get_json()
    notif_type = data.get("type")
    if not notif_type:
        return jsonify({"error": "Notification type is required."}), 400

    subject = data.get("subject", f"{notif_type} Notice")
    content = ""
    template_file = "Custom.docx"
    current_date = datetime.now().strftime("%d %B %Y")

    replacements = {
        "{SUBJECT}": subject,
        "{CURRENT_DATE}": current_date,
        "{SOURCE}": "PRINCIPAL",
        "{COPY_TO}": "All Departments HODs\nDean Office\nR&D\nStudents",
    }

    doc = None

    if notif_type == "Faculty Meeting":
        meeting_time = data.get("date_time")
        if not meeting_time:
            return jsonify({"error": "Meeting date/time is required."}), 400

        formatted_date = format_date(meeting_time, with_time=True)
        template_file = "Department_Meeting.docx"
        replacements["{MEETING_DATE}"] = formatted_date
        doc = Document(f"templates/{template_file}")

    elif notif_type in predefined_templates:
        # Class Schedule Changes, Exam Notifications, Result Announcements
        template_content = predefined_templates[notif_type]

        if notif_type == "Class Schedule Changes":
            start_date = data.get("start_date")
            if not start_date:
                return jsonify({"error": "Start date is required."}), 400
            content = template_content.replace("{{START_DATE}}", format_date(start_date))

        elif notif_type == "Exam Notifications":
            exam_type = data.get("exam_type")
            exam_date = data.get("exam_date")
            if not exam_type or not exam_date:
                return jsonify({"error": "Exam type and date are required."}), 400
            content = template_content.replace("{{EXAM_TYPE}}", exam_type)\
                                      .replace("{{EXAM_DATE}}", format_date(exam_date))

        elif notif_type == "Result Announcements":
            result_date = data.get("result_date")
            if not result_date:
                return jsonify({"error": "Result date is required."}), 400
            content = template_content.replace("{{RESULT_DATE}}", format_date(result_date))

        replacements["{CONTENT}"] = content
        doc = Document(f"templates/{template_file}")

    elif notif_type == "Custom":
        content = data.get("content", "").strip()
        if not subject or not content:
            return jsonify({"error": "Subject and content are required for custom notices."}), 400
        replacements["{CONTENT}"] = content
        doc = Document(f"templates/{template_file}")

    else:
        return jsonify({"error": f"Unsupported notification type: {notif_type}"}), 400

    # Replace placeholders
    for p in doc.paragraphs:
        for run in p.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Save and convert to PDF
    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').lower()}_{unique_id}"
    word_path = os.path.join(GENERATED_LETTERS_PATH, f"{filename_base}.docx")
    pdf_path = word_path.replace('.docx', '.pdf')

    doc.save(word_path)

    # Convert DOCX → PDF using Word automation
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        "message": "Notice generated successfully.",
        "download_link": f"http://localhost:5001/download/{os.path.basename(pdf_path)}"
    })


@bp.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(GENERATED_LETTERS_PATH, filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found."}), 404
    return send_from_directory(GENERATED_LETTERS_PATH, filename, as_attachment=True)