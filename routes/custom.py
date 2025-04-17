from flask import Blueprint, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pythoncom
import comtypes.client
import os
import re
import uuid

bp = Blueprint('custom', __name__)
GENERATED_LETTERS_PATH = 'generated_letters'
os.makedirs(GENERATED_LETTERS_PATH, exist_ok=True)

templates = {
    "Lost & Found Announcements": """
Dear Students,

A {{TYPE}} item has been reported: {{ITEM}}.

If this belongs to you or you have any information about it, please report to the Admin Office at the earliest.

Let’s help each other by maintaining a supportive and responsible campus environment.
""",
    "Environmental Initiatives": """
Dear Students,

We are excited to announce an upcoming event: {{EVENT}} scheduled on {{DATE}}.

We encourage everyone to join hands in this initiative and contribute to a greener and healthier planet.

Together, let’s make a positive impact!
""",
    "Alumni Interactions / Guest Visits": """
Dear Students,

We are pleased to inform you of an upcoming session – Alumni Interactions / Guest Visits scheduled on {{DATE}} at {{VENUE}}.

Such interactions offer valuable insights, guidance, and inspiration. All students are encouraged to attend.

Make the most of this opportunity!
"""
}

def format_date(date_str):
    dt = datetime.strptime(date_str, "%Y-%m-%dT%H:%M")
    day = dt.day
    suffix = 'th' if 4 <= day <= 20 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return dt.strftime(f"%d{suffix} %B %Y %I:%M %p")

@bp.route('/generate-custom-notice', methods=['POST'])
def generate_notice():
    data = request.get_json()
    notif_type = data.get('type')
    subject = data.get('subject') or notif_type
    item = data.get('item', '')
    event = data.get('event', '')
    date = data.get('date', '')
    venue = data.get('venue', '')
    custom_content = data.get('content', '')

    formatted_date = format_date(date) if date else ''
    current_date = datetime.now().strftime('%B %d, %Y')

    # Generate content
    if notif_type in templates:
        content = templates[notif_type].replace("{{TYPE}}", notif_type.split()[0]) \
                                       .replace("{{ITEM}}", item) \
                                       .replace("{{EVENT}}", event) \
                                       .replace("{{DATE}}", formatted_date) \
                                       .replace("{{VENUE}}", venue)
    else:
        # Custom "Others" path
        content = custom_content

    unique_id = uuid.uuid4().hex
    filename_base = f"{notif_type.replace(' ', '_').replace('/', '_').replace('&', '_').lower()}_{unique_id}"
    word_path = f"{GENERATED_LETTERS_PATH}/{filename_base}.docx"
    pdf_path = word_path.replace('.docx', '.pdf')

    doc = Document("templates/Custom.docx")
    replacements = {
        "{CURRENT_DATE}": current_date,
        "{SOURCE}": "PRINCIPAL",
        "{SUBJECT}": subject,
        "{CONTENT}": content,
        "{DATE}": formatted_date,
        "{COPY_TO}": "All HODs\nAdmin Office\nStudents"
    }

    for p in doc.paragraphs:
        for run in p.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = re.sub(re.escape(placeholder), value, run.text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = re.sub(re.escape(placeholder), value, run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(word_path)

    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(word_path))
    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    doc.Close()
    word.Quit()

    return jsonify({
        'message': 'Notice generated successfully!',
        'download_link': f'http://localhost:5001/download/{filename_base}.pdf'
    })

@bp.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(GENERATED_LETTERS_PATH, filename, as_attachment=True)